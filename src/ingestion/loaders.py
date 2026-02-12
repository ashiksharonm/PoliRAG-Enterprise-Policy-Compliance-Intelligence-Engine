"""Document loaders for various formats."""

import hashlib
import json
import mimetypes
from abc import ABC, abstractmethod
from datetime import datetime
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import Dict, List, Optional, Any
from uuid import uuid4

import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
from loguru import logger

from src.models import Document, DocumentFormat, Role


class DocumentLoader(ABC):
    """Abstract base class for document loaders."""

    @abstractmethod
    def load(self, file_path: Path, **metadata: Any) -> Document:
        """Load document from file path."""
        pass

    @staticmethod
    def compute_hash(content: str) -> str:
        """Compute SHA-256 hash of content."""
        return hashlib.sha256(content.encode("utf-8")).hexdigest()

    @staticmethod
    def get_file_size(file_path: Path) -> int:
        """Get file size in bytes."""
        return file_path.stat().st_size


class PDFLoader(DocumentLoader):
    """Load PDF documents."""

    def load(self, file_path: Path, **metadata: Any) -> Document:
        """Load PDF with table support using pdfplumber."""
        logger.info(f"Loading PDF: {file_path}")
        
        try:
            # Use pdfplumber for better table extraction
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        # Convert table to markdown-style format
                        table_text = self._table_to_text(table)
                        text_parts.append(table_text)
            
            content = "\n\n".join(text_parts)
            
            if not content.strip():
                # Fallback to PyPDF2
                logger.warning(f"pdfplumber extracted empty content, trying PyPDF2")
                content = self._load_with_pypdf2(file_path)
            
            return self._create_document(
                file_path=file_path,
                content=content,
                format=DocumentFormat.PDF,
                **metadata
            )
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise

    def _load_with_pypdf2(self, file_path: Path) -> str:
        """Fallback PDF loader using PyPDF2."""
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)

    @staticmethod
    def _table_to_text(table: List[List[str]]) -> str:
        """Convert table to text format."""
        if not table:
            return ""
        
        lines = []
        lines.append("\n[TABLE]")
        for row in table:
            cleaned_row = [str(cell).strip() if cell else "" for cell in row]
            lines.append(" | ".join(cleaned_row))
        lines.append("[/TABLE]\n")
        return "\n".join(lines)

    def _create_document(self, file_path: Path, content: str, format: DocumentFormat, **metadata: Any) -> Document:
        """Create Document object."""
        return Document(
            id=uuid4(),
            filename=file_path.name,
            format=format,
            content_hash=self.compute_hash(content),
            size_bytes=self.get_file_size(file_path),
            tenant_id=metadata.get("tenant_id", "default"),
            role_scope=metadata.get("role_scope", [Role.READ_ONLY]),
            source_path=str(file_path),
            metadata={"content": content, **metadata}
        )


class DOCXLoader(DocumentLoader):
    """Load DOCX documents."""

    def load(self, file_path: Path, **metadata: Any) -> Document:
        """Load DOCX preserving structure."""
        logger.info(f"Loading DOCX: {file_path}")
        
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Preserve heading styles
                    if paragraph.style.name.startswith("Heading"):
                        level = paragraph.style.name.replace("Heading ", "")
                        text = f"\n{'#' * int(level)} {text}\n"
                    text_parts.append(text)
            
            # Extract tables
            for table in doc.tables:
                table_text = self._table_to_text(table)
                text_parts.append(table_text)
            
            content = "\n\n".join(text_parts)
            
            return Document(
                id=uuid4(),
                filename=file_path.name,
                format=DocumentFormat.DOCX,
                content_hash=self.compute_hash(content),
                size_bytes=self.get_file_size(file_path),
                tenant_id=metadata.get("tenant_id", "default"),
                role_scope=metadata.get("role_scope", [Role.READ_ONLY]),
                source_path=str(file_path),
                metadata={"content": content, **metadata}
            )
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise

    @staticmethod
    def _table_to_text(table) -> str:
        """Convert DOCX table to text."""
        lines = ["\n[TABLE]"]
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
        lines.append("[/TABLE]\n")
        return "\n".join(lines)


class MarkdownLoader(DocumentLoader):
    """Load Markdown documents."""

    def load(self, file_path: Path, **metadata: Any) -> Document:
        """Load Markdown file."""
        logger.info(f"Loading Markdown: {file_path}")
        
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            return Document(
                id=uuid4(),
                filename=file_path.name,
                format=DocumentFormat.MARKDOWN,
                content_hash=self.compute_hash(content),
                size_bytes=self.get_file_size(file_path),
                tenant_id=metadata.get("tenant_id", "default"),
                role_scope=metadata.get("role_scope", [Role.READ_ONLY]),
                source_path=str(file_path),
                metadata={"content": content, **metadata}
            )
        except Exception as e:
            logger.error(f"Error loading Markdown {file_path}: {e}")
            raise


class JSONLoader(DocumentLoader):
    """Load JSON documents."""

    def load(self, file_path: Path, **metadata: Any) -> Document:
        """Load JSON file and convert to text."""
        logger.info(f"Loading JSON: {file_path}")
        
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            # Convert JSON to readable text format
            content = json.dumps(data, indent=2, ensure_ascii=False)
            
            return Document(
                id=uuid4(),
                filename=file_path.name,
                format=DocumentFormat.JSON,
                content_hash=self.compute_hash(content),
                size_bytes=self.get_file_size(file_path),
                tenant_id=metadata.get("tenant_id", "default"),
                role_scope=metadata.get("role_scope", [Role.READ_ONLY]),
                source_path=str(file_path),
                metadata={"content": content, "json_data": data, **metadata}
            )
        except Exception as e:
            logger.error(f"Error loading JSON {file_path}: {e}")
            raise


class EmailLoader(DocumentLoader):
    """Load Email (.eml) files."""

    def load(self, file_path: Path, **metadata: Any) -> Document:
        """Load EML file."""
        logger.info(f"Loading Email: {file_path}")
        
        try:
            with open(file_path, "rb") as f:
                msg = BytesParser(policy=policy.default).parse(f)
            
            # Extract email components
            subject = msg.get("subject", "")
            from_addr = msg.get("from", "")
            to_addr = msg.get("to", "")
            date = msg.get("date", "")
            
            # Extract body
            body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body += part.get_content()
            else:
                body = msg.get_content()
            
            # Format as readable text
            content = f"""Subject: {subject}
From: {from_addr}
To: {to_addr}
Date: {date}

{body}
"""
            
            return Document(
                id=uuid4(),
                filename=file_path.name,
                format=DocumentFormat.EMAIL,
                content_hash=self.compute_hash(content),
                size_bytes=self.get_file_size(file_path),
                tenant_id=metadata.get("tenant_id", "default"),
                role_scope=metadata.get("role_scope", [Role.READ_ONLY]),
                source_path=str(file_path),
                metadata={
                    "content": content,
                    "subject": subject,
                    "from": from_addr,
                    "to": to_addr,
                    "date": date,
                    **metadata
                }
            )
        except Exception as e:
            logger.error(f"Error loading Email {file_path}: {e}")
            raise


class TextLoader(DocumentLoader):
    """Load plain text documents."""

    def load(self, file_path: Path, **metadata: Any) -> Document:
        """Load text file."""
        logger.info(f"Loading Text: {file_path}")
        
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            return Document(
                id=uuid4(),
                filename=file_path.name,
                format=DocumentFormat.TEXT,
                content_hash=self.compute_hash(content),
                size_bytes=self.get_file_size(file_path),
                tenant_id=metadata.get("tenant_id", "default"),
                role_scope=metadata.get("role_scope", [Role.READ_ONLY]),
                source_path=str(file_path),
                metadata={"content": content, **metadata}
            )
        except Exception as e:
            logger.error(f"Error loading Text {file_path}: {e}")
            raise


class LoaderFactory:
    """Factory for creating document loaders."""

    _loaders = {
        ".pdf": PDFLoader,
        ".docx": DOCXLoader,
        ".md": MarkdownLoader,
        ".markdown": MarkdownLoader,
        ".json": JSONLoader,
        ".eml": EmailLoader,
        ".txt": TextLoader,
    }

    @classmethod
    def get_loader(cls, file_path: Path) -> DocumentLoader:
        """Get appropriate loader for file."""
        suffix = file_path.suffix.lower()
        
        if suffix not in cls._loaders:
            raise ValueError(f"Unsupported file format: {suffix}")
        
        loader_class = cls._loaders[suffix]
        return loader_class()

    @classmethod
    def supported_formats(cls) -> List[str]:
        """Get list of supported file formats."""
        return list(cls._loaders.keys())